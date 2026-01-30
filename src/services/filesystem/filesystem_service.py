import io
import os
from uuid import UUID, uuid4
import docx
from fastapi import UploadFile
from pypdf import PdfReader
import pandas as pd
from sqlalchemy import text

from src.models.documents.document import Document
from src.models.documents.document_chunk import DocumentChunk
from src.models.file_system.file_system_node import FileSystemNode

UPLOAD_DIR = "storage/documents"

class FileSystemService:
    def __init__(self, repo, embedding_model):
        self.__repo = repo
        self.__embedding_model = embedding_model
        # self.__embedding_model = OllamaEmbeddings(
        #     model="nomic-embed-text",
        #     base_url="http://localhost:11434"
        # )
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)

    async def __extract_text(self, file_content: bytes, filename: str) -> str:
        extension = filename.split('.')[-1].lower()
        try:
            if extension in ['txt', 'md', 'csv']:
                return file_content.decode('utf-8')
            elif extension == 'pdf':
                reader = PdfReader(io.BytesIO(file_content))
                return "".join([page.extract_text() for page in reader.pages])
            elif extension == 'docx':
                doc = docx.Document(io.BytesIO(file_content))
                return "\n".join([para.text for para in doc.paragraphs])
            elif extension in ['xlsx', 'xls']:
                df = pd.read_excel(io.BytesIO(file_content))
                return df.to_string()

            return "No content"
        except Exception as e:
            print(f"Extraction error for {extension}: {e}")
            return ""

    async def __chunk_and_save(self, session, doc_id: UUID, text: str):
        chunk_size = 800 # reduced chunk size for ollama nomic embeddings
        overlap = 150 # reduced overlap for ollama nomic embeddings
        chunks = []

        raw_text_chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            raw_text_chunks.append(text[i : i + chunk_size])

        if not raw_text_chunks:
            return
        # ollama embedding model
        embeddings = await self.__embedding_model.aembed_documents(raw_text_chunks)
        # hugging face embedding model

        for i, content in enumerate(raw_text_chunks):
            chunks.append(DocumentChunk(
                document_id=doc_id,
                chunk_index=i,
                text=content,
                embedding=embeddings[i]
            ))

        self.__repo.save_chunks(session, chunks)

    def get_similar_chunks(self, session, query_vector: list[float], limit: int = 5):
        return self.__repo.get_similar_chunks(session, query_vector, limit)

    def get_all_nodes(self, session):
        return self.__repo.get_all_nodes(session)

    async def create_node(self, session, name: str, type: str, parent_id: UUID, file: UploadFile = None):
        node = FileSystemNode(name=name, type=type, parent_id=parent_id)
        saved_node = self.__repo.save_node(session, node)

        if type == "file" and file:
            content = await file.read()

            file_extension = os.path.splitext(file.filename)[1]
            storage_path = os.path.join(UPLOAD_DIR, f"{uuid4()}{file_extension}")

            with open(storage_path, "wb") as buffer:
                buffer.write(content)

            document = Document(
                file_node_id=saved_node.id,
                # document_type=file.content_type,
                storage_path=storage_path
            )
            saved_doc = self.__repo.save_document(session, document)
            extracted_text = await self.__extract_text(content, file.filename)

            if extracted_text:
                await self.__chunk_and_save(session, saved_doc.id, extracted_text)

        return saved_node

    def get_node_and_all_descendants(self, session, node_id: UUID):
        # finds the parent and recursively finds all children/grandchildren
        query = text("""
            WITH RECURSIVE tree AS (
                SELECT * FROM file_nodes WHERE id = :node_id
                UNION ALL
                SELECT fn.* FROM file_nodes fn
                JOIN tree t ON fn.parent_id = t.id
            )
            SELECT * FROM tree;
        """)
        return session.query(FileSystemNode).from_statement(query).params(node_id=node_id).all()

    def delete_node(self, session, node_id: UUID):
        node = session.query(FileSystemNode).filter(FileSystemNode.id == node_id).first()
        if not node:
            return None

        # BEFORE the database delete to ensure data availability
        file_paths_to_remove = []
        if node.type == "file" and node.document:
            file_paths_to_remove.append(node.document.storage_path)

        self.__repo.delete_node(session, node_id)

        for path in file_paths_to_remove:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as e:
                print(f"Error removing file {path}: {e}")

        return node_id